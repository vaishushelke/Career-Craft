from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
import os
import sqlite3
import joblib
import io
from functools import wraps

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

app = Flask(__name__)
app.secret_key = os.urandom(24)

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'career_craft_v3.db')

def get_db_connection():
    conn = sqlite3.connect(DB_PATH, timeout=20)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Test integrity
        cursor.execute('PRAGMA integrity_check')
        if cursor.fetchone()[0] != 'ok':
             raise sqlite3.DatabaseError("Integrity check failed")

        # Users table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                email TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL
            )
        ''')

        # Profile Info table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS profile_info (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER UNIQUE,
                name TEXT, email TEXT, phone TEXT, phone2 TEXT, address TEXT,
                linkedin_url TEXT, github_url TEXT, portfolio_url TEXT, summary TEXT,
                theme_color TEXT DEFAULT '#06b6d4',
                title_font_size TEXT DEFAULT '24px',
                subtitle_font_size TEXT DEFAULT '18px',
                text_font_size TEXT DEFAULT '14px',
                font_family TEXT DEFAULT 'Inter',
                template_id INTEGER DEFAULT 1,
                FOREIGN KEY (user_id) REFERENCES users(id)
            )
        ''')

        # Relational tables
        cursor.execute('CREATE TABLE IF NOT EXISTS education (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, degree TEXT, institution TEXT, years TEXT, FOREIGN KEY (user_id) REFERENCES users(id))')
        cursor.execute('CREATE TABLE IF NOT EXISTS experience (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, company TEXT, position TEXT, years TEXT, description TEXT, FOREIGN KEY (user_id) REFERENCES users(id))')
        cursor.execute('CREATE TABLE IF NOT EXISTS skills (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, skill_name TEXT, skill_level TEXT DEFAULT "Intermediate", FOREIGN KEY (user_id) REFERENCES users(id))')
        cursor.execute('CREATE TABLE IF NOT EXISTS awards (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, title TEXT, FOREIGN KEY (user_id) REFERENCES users(id))')
        
        conn.commit()
        conn.close()
    except (sqlite3.DatabaseError, sqlite3.OperationalError) as e:
        print(f"Database error or corruption detected: {e}. Recreating...")
        if os.path.exists(DB_PATH):
            import time
            conn.close() if 'conn' in locals() else None
            try:
                os.remove(DB_PATH)
            except Exception:
                # If file is locked, just rename it
                os.rename(DB_PATH, f"{DB_PATH}.corrupted_{int(time.time())}")
        # Recursive call to recreate
        init_db()

init_db()

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please login first.', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if not username or not email or not password:
            flash('All fields are required.', 'danger')
            return redirect(url_for('register'))
            
        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return redirect(url_for('register'))
            
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            # Note: No hashing as per user request for simplicity
            cursor.execute('INSERT INTO users (username, email, password) VALUES (?, ?, ?)', (username, email, password))
            user_id = cursor.lastrowid
            # Create profile and update name with username
            cursor.execute('INSERT INTO profile_info (user_id, name, email) VALUES (?, ?, ?)', (user_id, username, email))
            conn.commit()
            conn.close()
            session['user_id'] = user_id
            session['username'] = username
            flash(f'Welcome, {username}! Account created.', 'success')
            return redirect(url_for('dashboard'))
        except sqlite3.IntegrityError:
            conn.close()
            flash('Username or Email already exists.', 'danger')
            
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        identifier = request.form.get('username') # acts as username or email
        password = request.form.get('password')
        
        conn = get_db_connection()
        user = conn.execute('SELECT * FROM users WHERE (username = ? OR email = ?) AND password = ?', (identifier, identifier, password)).fetchone()
        conn.close()
        
        if user:
            session['user_id'] = user['id']
            session['username'] = user['username']
            flash(f'Welcome back, {user["username"]}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid credentials.', 'danger')
            
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Logged out successfully.', 'info')
    return redirect(url_for('home'))

# Dashboard (Profile Info)
@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
def dashboard():
    conn = get_db_connection()
    cursor = conn.cursor()
    if request.method == 'POST':
        name = request.form.get('name', '')
        email = request.form.get('email', '')
        phone = request.form.get('phone', '')
        phone2 = request.form.get('phone2', '')
        address = request.form.get('address', '')
        linkedin_url = request.form.get('linkedin_url', '')
        github_url = request.form.get('github_url', '')
        portfolio_url = request.form.get('portfolio_url', '')
        summary = request.form.get('summary', '')
        theme_color = request.form.get('theme_color', '#06b6d4')
        title_size = request.form.get('title_font_size', '24px')
        subtitle_size = request.form.get('subtitle_font_size', '18px')
        text_size = request.form.get('text_font_size', '14px')

        cursor.execute('''
            UPDATE profile_info 
            SET name=?, email=?, phone=?, phone2=?, address=?, linkedin_url=?, github_url=?, portfolio_url=?, summary=?,
                theme_color=?, title_font_size=?, subtitle_font_size=?, text_font_size=?
            WHERE user_id=?
        ''', (name, email, phone, phone2, address, linkedin_url, github_url, portfolio_url, summary, 
              theme_color, title_size, subtitle_size, text_size, session['user_id']))
        conn.commit()
        flash('Profile updated!', 'success')
        
    cursor.execute('SELECT * FROM profile_info WHERE user_id=?', (session['user_id'],))
    profile = cursor.fetchone()
    conn.close()
    return render_template('dashboard.html', profile=profile)

@app.route('/builder/education', methods=['GET', 'POST'])
@login_required
def builder_education():
    conn = get_db_connection()
    cursor = conn.cursor()
    if request.method == 'POST':
        degree = request.form.get('degree')
        institution = request.form.get('institution')
        years = request.form.get('years')
        cursor.execute('INSERT INTO education (user_id, degree, institution, years) VALUES (?, ?, ?, ?)', (session['user_id'], degree, institution, years))
        conn.commit()
    
    cursor.execute('SELECT * FROM education WHERE user_id=?', (session['user_id'],))
    entries = cursor.fetchall()
    conn.close()
    return render_template('builder_education.html', entries=entries)

@app.route('/builder/education/delete/<int:id>')
@login_required
def delete_education(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM education WHERE id=? AND user_id=?', (id, session['user_id']))
    conn.commit()
    conn.close()
    return redirect(url_for('builder_education'))

@app.route('/builder/experience', methods=['GET', 'POST'])
@login_required
def builder_experience():
    conn = get_db_connection()
    cursor = conn.cursor()
    if request.method == 'POST':
        company = request.form.get('company')
        position = request.form.get('position')
        years = request.form.get('years')
        description = request.form.get('description')
        cursor.execute('INSERT INTO experience (user_id, company, position, years, description) VALUES (?, ?, ?, ?, ?)', (session['user_id'], company, position, years, description))
        conn.commit()
    
    cursor.execute('SELECT * FROM experience WHERE user_id=?', (session['user_id'],))
    entries = cursor.fetchall()
    conn.close()
    return render_template('builder_experience.html', entries=entries)

@app.route('/builder/experience/delete/<int:id>')
@login_required
def delete_experience(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM experience WHERE id=? AND user_id=?', (id, session['user_id']))
    conn.commit()
    conn.close()
    return redirect(url_for('builder_experience'))

@app.route('/builder/skills', methods=['GET', 'POST'])
@login_required
def builder_skills():
    conn = get_db_connection()
    cursor = conn.cursor()
    if request.method == 'POST':
        skill_name = request.form.get('skill_name')
        skill_level = request.form.get('skill_level', 'Intermediate')
        cursor.execute('INSERT INTO skills (user_id, skill_name, skill_level) VALUES (?, ?, ?)', (session['user_id'], skill_name, skill_level))
        conn.commit()
    
    cursor.execute('SELECT * FROM skills WHERE user_id=?', (session['user_id'],))
    entries = cursor.fetchall()
    conn.close()
    return render_template('builder_skills.html', entries=entries)

@app.route('/builder/skills/delete/<int:id>')
@login_required
def delete_skill(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM skills WHERE id=? AND user_id=?', (id, session['user_id']))
    conn.commit()
    conn.close()
    return redirect(url_for('builder_skills'))

@app.route('/builder/awards', methods=['GET', 'POST'])
@login_required
def builder_awards():
    conn = get_db_connection()
    cursor = conn.cursor()
    if request.method == 'POST':
        title = request.form.get('title')
        if title:
            cursor.execute('INSERT INTO awards (user_id, title) VALUES (?, ?)', (session['user_id'], title))
        
        font_size = request.form.get('font_size', '14px')
        font_family = request.form.get('font_family', 'Inter')
        template_id = request.form.get('template_id', 1)
        cursor.execute('UPDATE profile_info SET text_font_size=?, font_family=?, template_id=? WHERE user_id=?', (font_size, font_family, template_id, session['user_id']))
        conn.commit()
    
    cursor.execute('SELECT * FROM awards WHERE user_id=?', (session['user_id'],))
    entries = cursor.fetchall()
    cursor.execute('SELECT * FROM profile_info WHERE user_id=?', (session['user_id'],))
    profile = cursor.fetchone()
    conn.close()
    return render_template('builder_awards.html', entries=entries, profile=profile)

@app.route('/builder/awards/delete/<int:id>')
@login_required
def delete_award(id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM awards WHERE id=? AND user_id=?', (id, session['user_id']))
    conn.commit()
    conn.close()
    return redirect(url_for('builder_awards'))

@app.route('/resume')
@login_required
def resume():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM profile_info WHERE user_id=?', (session['user_id'],))
    profile = cursor.fetchone()
    
    if not profile:
        conn.close()
        flash('Please complete your profile first.', 'warning')
        return redirect(url_for('dashboard'))
        
    cursor.execute('SELECT * FROM education WHERE user_id=?', (session['user_id'],))
    education = cursor.fetchall()
    cursor.execute('SELECT * FROM experience WHERE user_id=?', (session['user_id'],))
    experience = cursor.fetchall()
    cursor.execute('SELECT * FROM skills WHERE user_id=?', (session['user_id'],))
    skills = cursor.fetchall()
    cursor.execute('SELECT * FROM awards WHERE user_id=?', (session['user_id'],))
    awards = cursor.fetchall()
    conn.close()
    
    user = dict(profile)
    user.update({
        'education_list': education,
        'experience_list': experience,
        'skills_list': skills,
        'awards_list': awards
    })
    return render_template('resume.html', user=user)

@app.route('/portfolio')
@login_required
def portfolio():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM profile_info WHERE user_id=?', (session['user_id'],))
    profile = cursor.fetchone()
    
    if not profile:
        conn.close()
        flash('Please complete your profile first.', 'warning')
        return redirect(url_for('dashboard'))
        
    cursor.execute('SELECT * FROM education WHERE user_id=?', (session['user_id'],))
    education = cursor.fetchall()
    cursor.execute('SELECT * FROM experience WHERE user_id=?', (session['user_id'],))
    experience = cursor.fetchall()
    cursor.execute('SELECT * FROM skills WHERE user_id=?', (session['user_id'],))
    skills = cursor.fetchall()
    cursor.execute('SELECT * FROM awards WHERE user_id=?', (session['user_id'],))
    awards = cursor.fetchall()
    conn.close()
    
    user = dict(profile)
    user.update({
        'education_list': education,
        'experience_list': experience,
        'skills_list': skills,
        'awards_list': awards
    })
    return render_template('portfolio.html', user=user)

@app.route('/recommend')
@login_required
def recommend():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT skill_name FROM skills WHERE user_id=?', (session['user_id'],))
    skills_data = cursor.fetchall()
    cursor.execute('SELECT * FROM profile_info WHERE user_id=?', (session['user_id'],))
    profile = cursor.fetchone()
    conn.close()
    
    if not profile:
        flash('Please complete your profile first.', 'warning')
        return redirect(url_for('dashboard'))
        
    if not skills_data:
        flash('Please add some skills first.', 'error')
        return redirect(url_for('builder_skills'))
        
    skills_str = ", ".join([s['skill_name'] for s in skills_data])
    
    model_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'model', 'career_model.pkl')
    roadmap_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'model', 'roadmaps.pkl')
    
    # Rule-based fallback for immediate functionality
    roadmaps_data = {
        "Front-end Developer": ["html", "css", "javascript", "react", "vue", "angular", "frontend", "typescript", "sass", "bootstrap", "tailwind"],
        "Back-end Developer": ["python", "flask", "django", "node.js", "express", "backend", "java", "spring", "sql", "postgresql", "mongodb", "apis"],
        "Data Scientist": ["python", "pandas", "numpy", "scikit-learn", "statistics", "data science", "machine learning", "visualization", "matplotlib", "seaborn", "r", "sql"],
        "AI/ML Engineer": ["machine learning", "ai", "artificial intelligence", "tensorflow", "pytorch", "deep learning", "nlp", "cv", "keras", "neural networks"],
        "Database Administrator": ["sql", "oracle", "mysql", "postgresql", "dba", "database administration", "management", "backup", "query optimization", "indexing"],
        "Mobile App Developer": ["android", "kotlin", "swift", "ios", "flutter", "mobile", "react native", "mobile development"],
        "DevOps Engineer": ["aws", "azure", "cloud", "devops", "docker", "kubernetes", "jenkins", "ci/cd", "linux", "sysadmin", "terraform"]
    }

    roadmap_steps_dict = {
        "Front-end Developer": "Learn HTML/CSS & Modern Layouts|Master JavaScript (ES6+)|Build Projects with React or Vue|Learn State Management (Redux/Pinia)|Understand Responsive Design & Web Accessibility",
        "Back-end Developer": "Master a Backend Language (Python/Node/Java)|Understand Relational Databases (SQL)|Design & Build RESTful APIs|Learn Authentication & Security|Deploy to Heroku/AWS/Vercel",
        "Data Scientist": "Master Python for Data Science (Pandas/NumPy)|Learn Exploratory Data Analysis (EDA)|Understand Machine Learning Algorithms|Master SQL for Data Extraction|Build a Portfolio of Real-world Projects",
        "AI/ML Engineer": "Master Linear Algebra & Calculus|Learn Key ML Libraries (Scikit-learn)|Deep Dive into TensorFlow or PyTorch|Understand NLP & Computer Vision|Optimize & Deploy Models to Production",
        "Database Administrator": "Master SQL & Relational Algebra|Learn Database Design & Normalization|Understand Indexing & Performance Tuning|Master Backup & Disaster Recovery|Learn Cloud Database Solutions (RDS/Cloud SQL)",
        "Mobile App Developer": "Learn Kotlin (Android) or Swift (iOS)|Understand UI/UX Design Patterns|Master State Management & API Integration|Learn Cross-platform Tools (Flutter/React Native)|Publish to Play Store/App Store",
        "DevOps Engineer": "Master Linux Command Line|Learn Containerization with Docker|Understand Orchestration with Kubernetes|Implement CI/CD Pipelines|Learn Infrastructure as Code (Terraform)",
        "Career Explorer": "Identify your core interests|Take introductory courses in various fields|Build small hobby projects|Connect with mentors in technical communities|Keep learning and building consistency"
    }

    try:
        if os.path.exists(model_path) and os.path.exists(roadmap_path):
            model = joblib.load(model_path)
            roadmaps_model = joblib.load(roadmap_path)
            predicted_career = model.predict([skills_str])[0]
            roadmap_str = roadmaps_model.get(predicted_career, "Career Explorer")
        else:
            # Simple keyword matching fallback
            skills_lower = skills_str.lower()
            scores = {career: sum(1 for kw in kws if kw in skills_lower) for career, kws in roadmaps_data.items()}
            predicted_career = max(scores, key=scores.get) if any(scores.values()) else "Career Explorer"
            roadmap_str = roadmap_steps_dict.get(predicted_career, "Career Explorer")
        
        roadmap_steps = roadmap_str.split('|')
    except Exception as e:
        predicted_career = "Career Explorer"
        roadmap_str = roadmap_steps_dict.get(predicted_career)
        roadmap_steps = roadmap_str.split('|')

    # Skill prediction logic
    ideal_skills = roadmaps_data.get(predicted_career, [])
    user_skills_lower = [s['skill_name'].lower().strip() for s in skills_data]
    suggested_skills = [s.title() for s in ideal_skills if s.lower() not in user_skills_lower][:5]

    return render_template('roadmap.html', 
                         career=predicted_career, 
                         steps=roadmap_steps, 
                         skills=skills_str, 
                         suggested_skills=suggested_skills,
                         user=dict(profile))

@app.route('/download_resume/word')
@login_required
def download_resume_word():
    if not HAS_DOCX:
        flash("Word export library missing.", "warning")
        return redirect(url_for('resume'))
    
    conn = get_db_connection()
    p = conn.execute('SELECT * FROM profile_info WHERE user_id=?', (session['user_id'],)).fetchone()
    edu = conn.execute('SELECT * FROM education WHERE user_id=?', (session['user_id'],)).fetchall()
    exp = conn.execute('SELECT * FROM experience WHERE user_id=?', (session['user_id'],)).fetchall()
    skills = conn.execute('SELECT * FROM skills WHERE user_id=?', (session['user_id'],)).fetchall()
    awards = conn.execute('SELECT * FROM awards WHERE user_id=?', (session['user_id'],)).fetchall()
    conn.close()

    if not p:
        return redirect(url_for('dashboard'))

    doc = Document()
    doc.add_heading(p['name'] or 'Your Name', 0)
    doc.add_paragraph(f"{p['email']} | {p['phone']} | {p['address']}")
    
    if p['summary']:
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(p['summary'])

    doc.add_heading('Experience', level=1)
    for e in exp:
        run = doc.add_paragraph(style='List Bullet').add_run(f"{e['position']} at {e['company']} ({e['years']})")
        run.bold = True
        doc.add_paragraph(e['description'])

    doc.add_heading('Education', level=1)
    for ed in edu:
        doc.add_paragraph(f"{ed['degree']} - {ed['institution']} ({ed['years']})", style='List Bullet')

    doc.add_heading('Skills', level=1)
    doc.add_paragraph(", ".join([s['skill_name'] for s in skills]))

    if awards:
        doc.add_heading('Awards', level=1)
        for a in awards:
            doc.add_paragraph(a['title'], style='List Bullet')

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True, download_name=f"{p['name'] or 'User'}_Resume.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_portfolio/word')
def download_portfolio_word():
    return download_resume_word()

if __name__ == '__main__':
    app.run(debug=True, port=8000)
